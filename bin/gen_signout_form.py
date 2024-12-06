#!/usr/bin/env python3

from docx import Document
from docx.shared import Pt, Mm, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.opc import constants
import pandas as pd
from pdf2image import convert_from_path
import argparse
import os
from datetime import datetime

parser = argparse.ArgumentParser(description='Generate Clinical Signout Form.')
parser.add_argument('-s', '--sample', type=str, dest='sample', required=True,
                    help='Name of a specimen')
parser.add_argument('-i', '--input_path', type=str, dest='input_path',
                    required=False, default="/mnt",
                    help='Path to SangerSeq analysis results')
parser.add_argument('-o', '--output_path', type=str, dest='output_path',
                    required=False, default="/mnt",
                    help='Directory where Signout form write to')
parser.add_argument('--share_point', type=str, dest='share_point', required=False,
                    default= 'https://sharepoint.com/sites/LaboratoryInformatics/Sanger',
                    help='Laboratory Informatics Share Point')

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    """Adds a page number to a word form."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_run = paragraph.add_run()
    font = page_run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

def add_checkbox_to_cell(cell):
    """Adds a checkbox content control to a table cell."""

    # Create a new paragraph in the cell
    paragraph = cell.paragraphs[0]

    # Create the checkbox element
    checkbox = OxmlElement('w:sdt')
    sdtpr = OxmlElement('w:sdtPr')
    checkbox.append(sdtpr)
    sdtcontent = OxmlElement('w:sdtContent')
    checkbox.append(sdtcontent)

    # Add the checkbox control to the content
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    r.append(rpr)
    t = OxmlElement('w:t')
    t.text = '☐ '  # Add a space to make the checkbox visible
    r.append(t)
    sdtcontent.append(r)

    # Add the checkbox to the paragraph
    paragraph._p.append(checkbox)

def set_table_border(table, insideV=False):
    """Sets a border style to a table."""
    borders = OxmlElement('w:tblBorders')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4')
    borders.append(bottom_border)
    h_border = OxmlElement('w:insideH')
    h_border.set(qn('w:val'), 'single')
    h_border.set(qn('w:sz'), '4')
    borders.append(h_border)
    if (insideV):
        v_border = OxmlElement('w:insideV')
        v_border.set(qn('w:val'), 'single')
        v_border.set(qn('w:sz'), '4')
        borders.append(v_border)
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '4')
    borders.append(top_border)
    table._tbl.tblPr.append(borders)

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph the hyperlink points to.
    :param url: The address of a web page or image
    :param text: The anchor text for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def change_table_border_color(table, color):
    """Sets a border color to a table."""

    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')

        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # Set border style if needed
            border.set(qn('w:sz'), '2')  # Set border size if needed
            border.set(qn('w:color'), color)  # Set the desired color
            tcBorders.append(border)

        tcPr.append(tcBorders)

def shade_table_header(table, color):
    """Shades the header row of a table in a document."""
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:shd')
        tcVAlign.set(ns.qn('w:fill'), color)
        tcPr.append(tcVAlign)

def add_table_footnote():
    """Add a table headnote. """
    list_paragraph = doc.add_paragraph()
    run = list_paragraph.add_run('\u2461' + 
        ' Read Length refers to the total number of base pairs (bp) sequenced from the forward ' +
        'or reverse sequences in a specimen.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color
    list_paragraph.add_run('\n')
    run = list_paragraph.add_run('\u2462' + 
        ' Phred scaled Quality Score (QS) is a measure of confidence assigned to a sequencing basecall.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color
    list_paragraph.add_run('\n')
    run = list_paragraph.add_run('\u2464' + 
        ' Contig Length refers to the length of a contiguous Sanger sequence generated by ' + 
        'assembling forward and reverse sequences in a specimen.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208)# Light blue color
    list_paragraph.add_run('\n')
    run = list_paragraph.add_run('\u2465' + 
        ' Signal Peak corresponds to a specific nucleotide base, with the peak height indicating ' +
        'the signal strength for that base.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color
    list_paragraph.add_run('\n')
    run = list_paragraph.add_run('\u2466' + 
        ' Mixed Base\u0025 refers to the percentage of mixed bases in forward or reverse sequences of a specimen.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color

def add_pherogram_note():
    """Add a table note for e-pherogram."""
    list_paragraph = doc.add_paragraph()
    run = list_paragraph.add_run(u'*' + 
        'A visual representation of DNA sequence data generated through Sanger sequencing, ' \
                'displayed as a graph where peaks on the y-axis represent the intensity of fluorescent '\
                'signals corresponding to different nucleotides (A, C, G, and T), '\
                    'with the position on the x-axis indicating the fragment length.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color

def add_contig_note():
    """Add a header note to the contig sequence table."""
    list_paragraph = doc.add_paragraph()
    run = list_paragraph.add_run(u'*' + 
        ' A Sanger sequencing consensus sequence is a calculated DNA sequence that represents the most frequently '\
            'occurring nucleotide at each position, derived by aligning and comparing forward and reverse '\
                'Sanger sequencing reads of a specimen.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color

def add_table_headnote():
    """Add a header note to a table."""
    list_paragraph = doc.add_paragraph()
    run = list_paragraph.add_run('\u2460' + 
        ' An unique identifier assigned to a sequence record within the NCBI Reference Sequence (RefSeq) database.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color
    list_paragraph.add_run('\n')
    run = list_paragraph.add_run('\u2461' + 
        ' The percentage of identical matches between the query and subject over the aligned region.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color
    list_paragraph.add_run('\n')
    run = list_paragraph.add_run('\u2462' + 
        ' Alignment score is based on the number of matches, mismatches, and gaps ' +
        'and indicates the sequence similarity between the query and subject.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208)# Light blue color
    list_paragraph.add_run('\n')
    run = list_paragraph.add_run('\u2464' + 
        ' Is an estimate of the Expected number of random alignments with a particular score or ' +
        'better that could be found by chance in each database search. E-values close to zero ' +
        'correspond to more significant alignments.')
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(4, 118, 208) # Light blue color

def create_seq_table(heading_txt, fasta_file):
    """Create a sequence table."""
    doc.add_heading(heading_txt, level=1)
    if 'Consensus' in heading_txt:
        add_contig_note()
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    change_table_border_color(table, 'D9D9D9') 
    table.width = Inches(6) 
    table.autofit = False
    table.allow_autofit = True

    with open(fasta_file) as f:
        table.cell(0, 0).text = f.readline().strip()
        for paragraph in table.cell(0, 0).paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    run.font.bold = False

def remove_table_border(table):
    """Remove borders from cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element.tcPr
            tc.left = None
            tc.top = None
            tc.right = None
            tc.bottom = None
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT

def set_table_font_name_size(table, font_name, font_size, align='center'):
    """Set a font name and size to a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if align=='left':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

args = parser.parse_args()

qc_csv = os.path.join(args.input_path, args.sample, "%s_qc_metrics.csv"%args.sample)
fwd_chromatogram = os.path.join(args.input_path, args.sample, "%s_F.chromatogram_100bases.pdf"%args.sample)
rev_chromatogram = os.path.join(args.input_path, args.sample, "%s_R.chromatogram_100bases.pdf"%args.sample)
fwd_chromatogram_pdf = "/".join([args.share_point, args.sample, "%s_F.chromatogram.pdf?csf=1&web=1"%args.sample])
rev_chromatogram_pdf = "/".join([args.share_point, args.sample, "%s_R.chromatogram.pdf?csf=1&web=1"%args.sample])
contig_fasta = os.path.join(args.input_path, args.sample, "%s_consensus_sequence.txt"%args.sample)
forward_fasta = os.path.join(args.input_path, args.sample, "%s_forward_sequence.txt"%args.sample)
reverse_fasta = os.path.join(args.input_path, args.sample, "%s_reverse_sequence.txt"%args.sample)
blast_tsv = os.path.join(args.input_path, args.sample, "blast", "%s_blast.tsv"%args.sample)
blast_html = "/".join([args.share_point, args.sample, "blast", "%s_blast_identity.html?csf=1&web=1"%args.sample])

# Create a new document
doc = Document()

# Add a title
title = doc.add_heading(args.sample, level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Format the datetime string
formatted_datetime = datetime.now().strftime("%Y-%m-%d %H:%M")
# Add a review table at the beginning of the document
review_tab = doc.add_table(rows=1, cols=2)

# Populate the table
review_tab.cell(0, 1).text = 'Date: ' + formatted_datetime
review_tab.cell(0, 0).text = 'Reviewed By: '
review_tab.autofit = False 
review_tab.allow_autofit = False
review_tab.cell(0, 1).width = Pt(130)
review_tab.cell(0, 0).width = Pt(300)
remove_table_border(review_tab)
set_table_font_name_size(review_tab, 'Calibri', 12, 'left')

# Add a section to the document
section = doc.sections[0]

# Add a header
section = doc.sections[0]
header = section.header
paragraph = header.paragraphs[0]
run = paragraph.add_run("Yale New Haven Hospital\u2120")
font = run.font
font.name = 'Calibri'
font.color.rgb = RGBColor(135, 206, 235) # Light blue color

doc.add_heading('Patient Demographic Information', level=1)
table1 = doc.add_table(rows=5, cols=2)
set_table_border(table1)
#table.style = 'Light Shading Accent 1'
table1.autofit = False
table1.allow_autofit = False
table1.cell(0, 0).text = 'Patient Name (MRN): '
table1.cell(1, 0).text = 'Specimen ID:'
table1.cell(2, 0).text = 'MPI #:'
table1.cell(3, 0).text = 'Specimen Source:'
table1.cell(4, 0).text = 'Organism Morphology:'
for row in table1.rows:
    for cell in row.cells:
        cell.width = Pt(120)
table1.cell(0, 1).width = Pt(300)
change_table_border_color(table1, 'D9D9D9')
set_table_font_name_size(table1, 'Calibri', 11)

doc.add_heading('Sample Quality Metrics', level=1)
qc_df = pd.read_csv(qc_csv)

# add table note
add_table_footnote()

table2 = doc.add_table(rows=qc_df.shape[0]+1, cols=qc_df.shape[1]+1)
qc_df.columns = ['Trace File Name','Read Len\u00B2','Mean QS\u00B3','QS20\u207A', 
                 'Contig Len\u2075','Signal Peak\u2076', 'Mixed Base%\u2077']
shade_table_header(table2, 'D9D9D9')
change_table_border_color(table2, 'D9D9D9')
table2.autofit = True
table2.allow_autofit = True

for row in table2.rows:
    for cell in row.cells:
        cell.width = Pt(45)
table2.cell(0, 0).width = Pt(100)
for i in range(qc_df.shape[1]):
    table2.cell(0, i).text = qc_df.columns[i]
table2.cell(0, qc_df.shape[1]).text = 'Status'

# add a data row for each item
for i, row in qc_df.iterrows():
    for j, col in enumerate(qc_df.columns):
        table2.cell(i+1, j).text = str(row[col])
    add_checkbox_to_cell(table2.cell(i+1,qc_df.shape[1]))
set_table_font_name_size(table2, 'Calibri', 10)

# Add an image
doc.add_heading('Patient Electropherogram\u002A', level=1)
add_pherogram_note()
fwd_h3 = doc.add_heading('1. Forward Electropherogram', level=2)
add_hyperlink(fwd_h3, fwd_chromatogram_pdf, '(View PDF)')

images = convert_from_path(fwd_chromatogram, poppler_path="/usr/bin")
for image in images:    
    image.save(os.path.join(args.output_path, args.sample, 'fwd_chromatogram.png'), 'PNG')
doc.add_picture(os.path.join(args.output_path, args.sample,'fwd_chromatogram.png'), width=Pt(500))
rev_h3 = doc.add_heading('2. Reverse Electropherogram', level=2)

add_hyperlink(rev_h3, rev_chromatogram_pdf, '(View PDF)')
images = convert_from_path(rev_chromatogram, poppler_path="/usr/bin")
for image in images:    
    image.save(os.path.join(args.output_path, args.sample,'rev_chromatogram.png'), 'PNG')
doc.add_picture(os.path.join(args.output_path, args.sample,'rev_chromatogram.png'), width=Pt(500))

# add consensus seqeuence tables
create_seq_table('Patient Consensus Sequence\u002A', contig_fasta)
create_seq_table('1. Forward Sanger Sequence', forward_fasta)
create_seq_table('2. Reverse Sanger Sequence', reverse_fasta)

try:
    blast_tab = doc.add_heading('BLAST Analysis of Patient Sequences', level=1)
    add_hyperlink(blast_tab, blast_html, " (View HTML Report)")
    blast_df = pd.read_csv(blast_tsv, header=None, comment="#", sep="\t", nrows=25)
    blast_df.columns = ['Accession#\u00B9', 'Identity%\u00B2', 'Score\u00B3', 'Length', 'Evalue\u2075', 'Subject']
    blast_df['Evalue\u2075'] = blast_df['Evalue\u2075'].apply(lambda x: round(x, 3))

    add_table_headnote()
    table4 = doc.add_table(rows=blast_df.shape[0]+1, cols=blast_df.shape[1])
    table4.style = 'Light Shading Accent 1'
    table4.autofit = False
    table4.allow_autofit = False
    for row in table4.rows:
        for cell in row.cells:
            cell.width = Pt(50)
    table4.cell(0, 0).width = Pt(80)
    table4.cell(0, 1).width = Pt(80)
    table4.cell(0, 5).width = Pt(160)
    for i in range(blast_df.shape[1]):
        table4.cell(0, i).text = blast_df.columns[i]

    # add a data row for each item
    for i, row in blast_df.iterrows():
        for j, col in enumerate(blast_df.columns):
            table4.cell(i+1, j).text = str(row[col])

    set_table_font_name_size(table4, 'Calibri', 10)
except:
    pass

doc.add_heading('Patient Signout', level=1)
table5 = doc.add_table(rows=1, cols=1)
table5.style = 'Table Grid'
change_table_border_color(table5, 'D9D9D9') 
for row in table5.rows:
    row.height = Cm(16)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.text = ' '
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

# Add page numbers to each section's footer
for section in doc.sections:
    add_page_number(section.footer.paragraphs[0])

# Save the document
doc.save('./%s_signout_form.docx'%args.sample)